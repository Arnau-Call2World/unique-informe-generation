from docx import Document
import pandas as pd

def insertar_tabla_categorias(doc: Document, df: pd.DataFrame):
    doc.add_paragraph("Tabla resumen por categorías:")

    columnas = ["Categoría", "Recibidas", "Atendidas_num", "Atendidas_%"]
    columnas_titulo = ["Categoría", "Recibidas | Número", "Atendidas | Número", "Atendidas | %"]

    table = doc.add_table(rows=1, cols=len(columnas))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columnas_titulo):
        hdr_cells[i].text = col

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(columnas):
            if col == "Atendidas_%":
                row_cells[i].text = f"{row[col]:.2f} %"
            else:
                row_cells[i].text = str(row[col])

def insertar_tabla_dias(doc: Document, df: pd.DataFrame):

    columnas = ["Día", "Categoría", "Recibidas", "Atendidas_num", "Atendidas_%"]
    columnas_titulo = ["Día", "Categoría", "Recibidas | Número", "Atendidas | Número", "Atendidas | %"]

    # Filtrar solo rojas
    df_filtrado = df.iloc[:-1]
    df_filtrado = df_filtrado[(df_filtrado["Atendidas_%"] < 85.0) & (df_filtrado["Atendidas_%"] > 0)]


    # Normalizar fecha como texto tipo dd/mm/yy
    df = df_filtrado.copy()
    df["Categoría"] = df["Categoría"].apply(lambda x: pd.to_datetime(x).strftime("%d/%m/%y") if pd.notnull(x) else "")

    table = doc.add_table(rows=1, cols=len(columnas))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columnas_titulo):
        hdr_cells[i].text = col

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row["Día"]  # Día de la semana
        row_cells[1].text = row["Categoría"]  # Fecha formateada
        row_cells[2].text = str(row["Recibidas"])
        row_cells[3].text = str(row["Atendidas_num"])
        row_cells[4].text = f"{row['Atendidas_%']:.2f} %"

def insertar_tabla_franjas(doc: Document, df: pd.DataFrame):

    columnas = ["Categoría", "Recibidas", "Atendidas_num", "Atendidas_%"]
    columnas_titulo = ["Categoría", "Recibidas | Número", "Atendidas | Número", "Atendidas | %"]

    # Filtrar solo filas rojas
    df_filtrado = df.iloc[:-1]
    df_filtrado = df_filtrado[(df_filtrado["Atendidas_%"] < 85.0) & (df_filtrado["Atendidas_%"] > 0)]

    table = doc.add_table(rows=1, cols=len(columnas))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columnas_titulo):
        hdr_cells[i].text = col

    for _, row in df_filtrado.iterrows():  # ← Aquí el cambio
        row_cells = table.add_row().cells
        for i, col in enumerate(columnas):
            if col == "Atendidas_%":
                row_cells[i].text = f"{row[col]:.2f} %"
            else:
                row_cells[i].text = str(row[col])
