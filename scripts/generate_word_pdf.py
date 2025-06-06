import os 
import sys
from io import BytesIO

# Añadir el directorio padre al path para importar correctamente
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from scripts.generate_variables import generar_variables_informe
from scripts.insert_tables import (
    insertar_tabla_categorias,
    insertar_tabla_dias,
    insertar_tabla_franjas,
)
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import os

def generar_correo_informe(
    mes: str,
    año: int,
    porcentaje_mes: float,
    porcentaje_mes_anterior: float,
    cumple_objetivo: bool,
    total_recibidas: int,
    total_no_atendidas: int,
    desbordadas: int,
    abandonadas: int,
    categorias_conflictivas: list,  # Lista de tuplas (nombre, porcentaje)
    tabla_categorias: str,
    resumen_dias_flojos: str,
    resumen_franjas_flojas: str,
) -> str:
    cumplimiento = "El porcentaje de atención al cliente es bajo, pues sigue sin llegar al mínimo aceptable de 85%." \
        if not cumple_objetivo else "El porcentaje de atención cumple con el objetivo mínimo del 85%."

    texto = f"""
El nivel de atención ha {'aumentado' if porcentaje_mes > porcentaje_mes_anterior else 'disminuido'} con un {porcentaje_mes:.2f}% versus al {porcentaje_mes_anterior:.2f}% del mes de {mes_anterior(mes)} de {año}. {cumplimiento}

Para visualizar mejor la evolución adjuntamos histórico de nivel de atención de lo que llevamos del año:

[GRÁFICO]

Informe de Colas:

Nos refleja que se han recibido un total de {total_recibidas} llamadas y no se han atendido {total_no_atendidas}.

De las {total_no_atendidas} llamadas que no han sido atendidas:

{desbordadas} llamadas han sido desbordadas por tiempo.
{abandonadas} llamadas han sido abandonadas.

Las colas que más han afectado al porcentaje de atención:
"""

    for cat, pct in categorias_conflictivas:
        texto += f"\n- {cat} con un {pct:.2f} %."

    texto += f"""

Copio cuadro para mayor visibilidad:

{tabla_categorias}

El informe por días, nos indica que los días más flojos de {mes.lower()}:

{resumen_dias_flojos}

El informe por franja horaria nos indica que las más afectadas son:

{resumen_franjas_flojas}
"""
    return texto

def mes_anterior(mes_actual: str) -> str:
    meses = [
        "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
        "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"
    ]
    idx = meses.index(mes_actual.capitalize())
    return meses[idx - 1] if idx > 0 else meses[-1]

import tempfile

def exportar_informe_word_y_pdf_por_bloques(
    texto_informe: str,
    tabla_categorias: pd.DataFrame,
    tabla_dias: pd.DataFrame,
    tabla_franjas: pd.DataFrame,
    fig: plt.Figure = None
) -> tuple[BytesIO, BytesIO]:
    
    doc = Document()
    doc.add_heading("Informe de llamadas", level=1)
    lineas = texto_informe.strip().split("\n")

    temp_img_path = None
    if fig is not None:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmpfile:
            temp_img_path = tmpfile.name
            fig.savefig(temp_img_path, bbox_inches='tight')

    for linea in lineas:
        linea = linea.strip()
        if not linea:
            continue

        if "[GRÁFICO]" in linea and temp_img_path:
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(temp_img_path, width=Inches(5.5))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif "[TABLA_CATEGORIAS]" in linea:
            insertar_tabla_categorias(doc, tabla_categorias)

        elif "[TABLA_DIAS]" in linea:
            insertar_tabla_dias(doc, tabla_dias)

        elif "[TABLA_FRANJAS]" in linea:
            insertar_tabla_franjas(doc, tabla_franjas)

        else:
            doc.add_paragraph(linea)

    # Guardar Word en archivo temporal
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_word:
        doc.save(tmp_word.name)
        tmp_word_path = tmp_word.name


    # Cargar Word y PDF en memoria
    with open(tmp_word_path, "rb") as f:
        word_bytes = BytesIO(f.read())


    # Eliminar temporales
    os.remove(tmp_word_path)
    if temp_img_path and os.path.exists(temp_img_path):
        os.remove(temp_img_path)

    # Retroceder punteros
    word_bytes.seek(0)


    return word_bytes

def main(mes_actual, categorias_csv, dias_csv, franjas_csv, historico_xlsx):
    meses_hasta_ahora = ["ENERO", "FEBRERO", "MARZO", "ABRIL", "MAYO"]
    path_paquete = [categorias_csv, dias_csv, franjas_csv, historico_xlsx]
    variables_informe, fig, tablas = generar_variables_informe(
        path_paquete,
        meses_hasta_ahora,
        mes_actual
    )

    cuerpo = generar_correo_informe(
        mes=mes_actual,
        año=2025,
        porcentaje_mes=variables_informe[0],
        porcentaje_mes_anterior=variables_informe[1],
        cumple_objetivo=variables_informe[2],
        total_recibidas=variables_informe[3],
        total_no_atendidas=variables_informe[4],
        desbordadas=variables_informe[5],
        abandonadas=variables_informe[6],
        categorias_conflictivas=variables_informe[7],
        tabla_categorias="[TABLA_CATEGORIAS]",
        resumen_dias_flojos="[TABLA_DIAS]",
        resumen_franjas_flojas="[TABLA_FRANJAS]"
    )
    word_path, pdf_path = exportar_informe_word_y_pdf_por_bloques(
        texto_informe=cuerpo,
        tabla_categorias=tablas['categorias'],
        tabla_dias=tablas['dias'],
        tabla_franjas=tablas['franjas'],
        fig=fig
    )
    return word_path, pdf_path



if __name__ == "__main__":
    MESES = [ "ENERO", "FEBRERO", "MARZO", "ABRIL", "MAYO", "JUNIO",
                       "JULIO", "AGOSTO", "SEPTIEMBRE", "OCTUBRE", "NOVIEMBRE", "DICIEMBRE" ]
    mes_actual = "MAYO"
    path_paquete = "data/paquete-mayo/"
    meses_hasta_ahora = MESES[:MESES.index(mes_actual)+1]

    variables_informe, fig, tablas = generar_variables_informe(path_paquete, meses_hasta_ahora, mes_actual)
    cuerpo = generar_correo_informe(
        mes=mes_actual,
        año=2025,
        porcentaje_mes=variables_informe[0],
        porcentaje_mes_anterior=variables_informe[1],
        cumple_objetivo=variables_informe[2],
        total_recibidas=variables_informe[3],
        total_no_atendidas=variables_informe[4],
        desbordadas=variables_informe[5],
        abandonadas=variables_informe[6],
        categorias_conflictivas=variables_informe[7],
        tabla_categorias="[TABLA_CATEGORIAS]",
        resumen_dias_flojos="[TABLA_DIAS]",
        resumen_franjas_flojas="[TABLA_FRANJAS]"
    )

    word_path, pdf_path = exportar_informe_word_y_pdf_por_bloques(
        texto_informe=cuerpo,
        path_word=f"Informe_Autoclima_2025.docx",
        tabla_categorias=tablas['categorias'],
        tabla_dias=tablas['dias'],
        tabla_franjas=tablas['franjas'],
        fig=fig
    )
